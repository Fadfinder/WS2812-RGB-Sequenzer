from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "WS2812_RGB_Sequenzer_Benutzerhandbuch.docx"
WIRING_PNG = ROOT / "docs" / "wemos_d1_mini_ws2812b_wiring.png"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 99, 112)
RED = RGBColor(155, 28, 28)


def set_font(run, size=11, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def heading(doc, text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def numbered(doc, text):
    return doc.add_paragraph(text, style="List Number")


def note(doc, title, text, danger=False):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, "FDECEC" if danger else "E8EEF5")
    set_cell_margins(cell, 140, 160, 140, 160)
    p = cell.paragraphs[0]
    r = p.add_run(f"{title}: ")
    set_font(r, bold=True, color=RED if danger else DARK_BLUE)
    set_font(p.add_run(text))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_two_col_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    widths = [Inches(1.8), Inches(4.7)]
    for i, label in enumerate(("Element", "Beschreibung")):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        set_cell_shading(cell, "E8EEF5")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        r = cell.paragraphs[0].add_run(label)
        set_font(r, bold=True, color=DARK_BLUE)
    for left, right in rows:
        cells = table.add_row().cells
        for i, value in enumerate((left, right)):
            cells[i].width = widths[i]
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_font(cells[i].paragraphs[0].add_run(value))
    return table


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.49)
    sec.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        styles[name].font.name = "Calibri"
        styles[name].font.size = Pt(11)
        styles[name].paragraph_format.space_after = Pt(4)
        styles[name].paragraph_format.line_spacing = 1.25

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("WS2812 RGB Sequenzer | Benutzerhandbuch"), 9, color=MUTED)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.add_run("Seite "), 9, color=MUTED)
    add_field(footer, "PAGE")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(86)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("BENUTZERHANDBUCH"), 11, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_font(p.add_run("WS2812 RGB Sequenzer"), 30, bold=True, color=DARK_BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(44)
    set_font(p.add_run("Wemos D1 mini, WS2812B und eigenständige Websteuerung"), 15, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("Einrichten, bedienen, Sequenzen erstellen und Fehler beheben"), 11, italic=True, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    set_font(p.add_run("Stand: 12. Juni 2026"), 10, color=MUTED)
    doc.add_page_break()

    heading(doc, "1. Überblick")
    body(doc, "Der WS2812 RGB Sequenzer steuert bis zu 50 adressierbare WS2812B-LEDs. Die Einstellungen werden über eine Weboberfläche vorgenommen und dauerhaft im Wemos gespeichert. Nach dem Speichern arbeitet das Gerät ohne Browser selbstständig weiter und startet die gespeicherten Sequenzen nach einem Stromausfall automatisch erneut.")
    add_two_col_table(doc, [
        ("Controller", "Wemos D1 mini / ESP8266"),
        ("LED-Typ", "WS2812B, 5 V, Datenanschluss DIN"),
        ("Maximale Anzahl", "50 LEDs"),
        ("Datenpin", "D4 am Wemos"),
        ("Speicherung", "Dauerhaft im LittleFS-Flash des Wemos"),
        ("Bedienung", "Eigenes WLAN und Browseroberfläche"),
    ])

    heading(doc, "2. Sicherheit und Stromversorgung")
    note(doc, "Wichtig", "Den Stepdown vor dem Anschließen des Wemos und des LED-Streifens mit einem Multimeter auf etwa 5,0 V einstellen. Mehr als 5 V kann LEDs und Wemos zerstören.", True)
    bullet(doc, "12-V-Netzteil nur am Eingang des Stepdown-Wandlers anschließen.")
    bullet(doc, "Wemos und WS2812B-Streifen gemeinsam mit 5 V aus dem Stepdown versorgen.")
    bullet(doc, "Alle GND-Anschlüsse müssen miteinander verbunden sein.")
    bullet(doc, "Den 1200-µF-Elko polrichtig anschließen: Minus-Markierung an GND, Plus an 5 V.")
    bullet(doc, "Bei voller weißer Helligkeit mit bis zu etwa 60 mA pro LED rechnen. Für 20 LEDs plus Wemos ist ein 5-V-/3-A-Stepdown geeignet.")
    bullet(doc, "Vor Arbeiten an der Verdrahtung immer die Stromversorgung trennen.")

    heading(doc, "3. Verdrahtung")
    if WIRING_PNG.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(WIRING_PNG), width=Inches(6.4))
        p.paragraph_format.keep_with_next = True
    body(doc, "Minimalanschluss:")
    add_two_col_table(doc, [
        ("12 V + / -", "An Stepdown IN+ / IN-"),
        ("Stepdown 5 V", "An Wemos 5V und WS2812B 5V"),
        ("Stepdown GND", "An Wemos G und WS2812B GND"),
        ("Wemos D4", "Über 330-Ohm-Widerstand an DIN des ersten WS2812B-Pixels"),
        ("1200-µF-Elko", "Plus an 5 V, Minus-Markierung an GND, nahe am Streifenanfang"),
    ])
    note(doc, "Hinweis", "Bei längeren Leitungen oder unzuverlässigem Datenempfang ist ein 74AHCT125 oder 74HCT125 als 3,3-V-auf-5-V-Pegelwandler empfehlenswert.")

    heading(doc, "4. Erststart und Verbindung")
    numbered(doc, "Wemos und LED-Streifen einschalten.")
    numbered(doc, "Am Smartphone, Tablet oder Computer das WLAN RGB-Sequenzer auswählen.")
    numbered(doc, "Das Standardkennwort 12345678 eingeben.")
    numbered(doc, "Falls sich die Bedienoberfläche nicht automatisch öffnet, im Browser http://rgb-sequenzer.local aufrufen.")
    numbered(doc, "Falls diese Adresse nicht funktioniert, die im Projekt verwendete Geräte-IP aufrufen.")
    note(doc, "WLAN", "Das Wemos-WLAN hat absichtlich keinen Internetzugang. Manche Geräte melden deshalb 'Kein Internet'. Die Verbindung trotzdem beibehalten.")

    heading(doc, "5. Grundbedienung")
    add_two_col_table(doc, [
        ("Abspielen", "Startet ab dem aktuell ausgewählten Schritt und spielt danach die Sequenzen normal weiter."),
        ("Stop", "Beendet die automatische Wiedergabe und aktiviert die direkte Live-Bearbeitung."),
        ("Speichern", "Speichert LED-Anzahl, WLAN-Daten, Gruppen, Sequenzen, Schritte und Aktivierungsstatus dauerhaft."),
        ("LED-Anzahl", "An die tatsächlich angeschlossene Zahl anpassen, maximal 50."),
    ])
    body(doc, "Die Schaltflächen Abspielen, Stop und Speichern bleiben beim Scrollen am oberen Bildschirmrand sichtbar, auch auf dem Smartphone.")

    heading(doc, "6. Sequenzen verwalten")
    bullet(doc, "Eine Sequenz enthält beliebig viele Schritte und besitzt einen Namen.")
    bullet(doc, "Sequenzen können aktiviert, deaktiviert, verschoben, eingefügt und gelöscht werden.")
    bullet(doc, "Deaktivierte Sequenzen werden beim Abspielen übersprungen; der Zustand bleibt nach einem Neustart erhalten.")
    bullet(doc, "Bei 'Dauer aus Schritten berechnen' wird die Sequenzdauer automatisch aus allen Schritten gebildet und bei jeder Änderung sofort aktualisiert.")
    bullet(doc, "Ohne automatische Dauer kann eine feste Gesamtdauer vorgegeben werden.")
    note(doc, "Zufallsdauer", "Hat ein Schritt eine zufällige Dauer von 0 bis X Sekunden, verwendet die automatische Sequenzdauer X als planbare Obergrenze.")

    heading(doc, "7. Schritte bearbeiten")
    numbered(doc, "Sequenz auswählen und einen vorhandenen Schritt anklicken oder einen neuen Schritt anlegen.")
    numbered(doc, "Schritt benennen und Dauer in Sekunden einstellen. Die Dauer kann auch direkt in der unteren Schrittliste geändert werden.")
    numbered(doc, "Optional eine Zufallsdauer in 0,1-Sekunden-Schritten festlegen.")
    numbered(doc, "Schritthelligkeit einstellen. Sie wirkt proportional auf die Helligkeiten der einzelnen LEDs und Effekte.")
    numbered(doc, "Farben und Effekte zuweisen. Beim Wechsel zu einem anderen Schritt werden Änderungen automatisch in den Arbeitsspeicher übernommen.")
    numbered(doc, "Mit Speichern das gesamte Projekt dauerhaft in den Wemos schreiben.")
    body(doc, "Im Stop-Modus zeigt der echte LED-Streifen den aktuell ausgewählten Schritt. Farb-, Helligkeits- und Effektänderungen erscheinen direkt auf den LEDs.")

    heading(doc, "8. Farben, Helligkeit und LED-Gruppen")
    bullet(doc, "Farbe und Helligkeit auswählen und danach die gewünschten LED-Schaltflächen anklicken.")
    bullet(doc, "Ein erneuter Klick mit derselben Farbe und Helligkeit schaltet die LED aus.")
    bullet(doc, "Jede LED kann innerhalb desselben Schritts eine eigene Farbe und Helligkeit besitzen.")
    bullet(doc, "Häufig verwendete LED-Auswahlen können als Gruppe gespeichert und benannt werden, zum Beispiel Fenster, Fahrzeug oder Bühne.")
    bullet(doc, "Gruppennamen stehen anschließend in der Effektzuordnung zur Auswahl.")
    body(doc, "Beispiel für proportionale Helligkeit: Eine LED mit 60 % Helligkeit und ein Schritt mit 50 % ergeben effektiv 30 %.")

    heading(doc, "9. Effekte")
    body(doc, "Mehrere Effekte können gleichzeitig in einem Schritt verwendet werden. Jeder Effekt besitzt eine eigene LED- oder Gruppenzuordnung, Geschwindigkeit und Helligkeit. Bei einfarbigen Effekten kann zusätzlich eine Effektfarbe gewählt werden.")
    add_two_col_table(doc, [
        ("Geschwindigkeit 0", "Sehr langsam; bei geeigneten Effekten bis ungefähr fünf Minuten."),
        ("Geschwindigkeit 50", "Standardwert; der Regenbogen benötigt ungefähr zehn Sekunden pro Durchlauf."),
        ("Geschwindigkeit 100", "Sehr schnell."),
        ("Fade-In / Fade-Out", "Laufen jeweils über die gesamte eingestellte Schrittdauer."),
        ("Farbübergang", "Übergang von der Grundfarbe zur Effektfarbe; ohne Grundfarbe beginnt er bei Schwarz."),
    ])
    body(doc, "Enthalten sind unter anderem Feuerflackern, Gewitter, Regenbogen, Schweißerlicht, Fotografen-Blitz, Polizei-Blitzer, Funkeln, Komet, Theater-Lauflicht, Pulsieren, Atmen, Lava, Zufallseffekte, Meteor, Sternfunkeln, Kerze, Sonnenaufgang, Sonnenuntergang, Scanner, Konfetti, Aurora, Giftgrün, Herzschlag, Fade-In, Fade-Out, Farbübergang, Blinklicht und Lauflicht.")

    heading(doc, "10. Speichern und automatischer Betrieb")
    bullet(doc, "Speichern schreibt das gesamte Projekt dauerhaft in den Flash-Speicher.")
    bullet(doc, "Nach Stromausfall oder Neustart startet die gespeicherte Wiedergabe automatisch.")
    bullet(doc, "Die Browseroberfläche liest beim Öffnen den aktuellen Datenstand aus dem Wemos.")
    bullet(doc, "Die Demo_Sequenz wird durch diese Firmware beim Hochladen kontrolliert neu angelegt, damit keine mehrfach angehängten Demo-Schritte entstehen.")
    note(doc, "Praxis", "Nach größeren Änderungen einmal Speichern drücken und anschließend den Wemos testweise aus- und wieder einschalten. So lässt sich die dauerhafte Speicherung direkt kontrollieren.")

    heading(doc, "11. WLAN-Name und Kennwort ändern")
    numbered(doc, "Neuen WLAN-Namen eingeben und den zugehörigen Übernehmen-Button drücken.")
    numbered(doc, "Neues Kennwort zweimal sichtbar eingeben. Es muss mindestens acht Zeichen lang sein.")
    numbered(doc, "Kennwort übernehmen und anschließend Speichern drücken.")
    numbered(doc, "Nach dem Neustart mit dem neuen WLAN-Namen und Kennwort verbinden.")
    note(doc, "Wichtig", "Die neuen Zugangsdaten sicher notieren. Nach der Änderung trennt sich die bisherige WLAN-Verbindung erwartungsgemäß.")

    heading(doc, "12. Fehlersuche")
    add_two_col_table(doc, [
        ("Keine LED leuchtet", "DIN statt DOUT verwenden, D4 prüfen, gemeinsame Masse prüfen, LED-Anzahl kontrollieren und 5 V messen."),
        ("Eine LED funktioniert, Streifen nicht", "Richtungspfeile/DIN prüfen, Lötstellen und erste LED kontrollieren, Leitung kürzen und Pegelwandler testen."),
        ("Falsche Farben / Flackern", "GND-Verbindung, 330-Ohm-Widerstand, 5-V-Spannung unter Last und Datenleitung prüfen."),
        ("Wemos-WLAN verschwindet", "5-V-Versorgung und Elko prüfen, LED-Helligkeit reduzieren, Stepdown-Stromreserve kontrollieren und Wemos neu starten."),
        ("Weiße Browserseite", "Seite neu laden, WLAN-Verbindung beibehalten, andere Adresse testen und Browser-Cache leeren."),
        ("Änderungen nach Neustart weg", "Nach Änderungen Speichern drücken und Bestätigung abwarten."),
        ("Sequenzdauer falsch", "Dauer aus Schritten aktivieren; die Anzeige muss sich beim Ändern einer Schrittdauer sofort aktualisieren."),
    ])

    heading(doc, "13. Technische Kurzreferenz")
    add_two_col_table(doc, [
        ("Board in Arduino IDE", "LOLIN(WEMOS) D1 R2 & mini"),
        ("Bibliothek", "Adafruit NeoPixel"),
        ("Aktueller Sketch", "Wemos_D1_Mini_RGB_Sequencer_Brightness.ino"),
        ("Datenpin", "D4"),
        ("LED-Spannung", "5 V"),
        ("Maximale LEDs", "50"),
        ("Standard-WLAN", "RGB-Sequenzer"),
        ("Standardkennwort", "12345678"),
    ])

    doc.core_properties.title = "WS2812 RGB Sequenzer - Benutzerhandbuch"
    doc.core_properties.subject = "Bedienung des Wemos D1 mini WS2812B Sequenzers"
    doc.core_properties.author = "WS2812 RGB Sequenzer Projekt"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
